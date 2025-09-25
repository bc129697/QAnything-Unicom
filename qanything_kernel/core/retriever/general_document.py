from qanything_kernel.utils.general_utils import get_time, get_table_infos, \
    html_to_markdown, clear_string, get_time_async, remove_think_tags, get_docx_toc
from qanything_kernel.utils.model_utils import num_tokens_embed
from typing import List, Optional
from qanything_kernel.configs.model_config import UPLOAD_ROOT_PATH, LOCAL_OCR_SERVICE_URL, IMAGES_ROOT_PATH, \
    DEFAULT_CHILD_CHUNK_SIZE, LOCAL_PDF_PARSER_SERVICE_URL, OUTLINE_EXTRACT_SYSTEM_PROMPT, SUMMARY_EXTRACT_SYSTEM_PROMPT
from qanything_kernel.configs.model_config import LLM_BASE_URL, LLM_API_KEY, LLM_MODEL_NAME, LLM_MAX_LENGTH, LLM_TOP_P, LLM_TEMPERATURE, LLM_MAX_OUTPUT_LENGTH
from langchain.docstore.document import Document
from qanything_kernel.utils.loader.my_recursive_url_loader import MyRecursiveUrlLoader
from qanything_kernel.utils.custom_log import insert_logger
from langchain_community.document_loaders import UnstructuredFileLoader, TextLoader
# from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_community.document_loaders import UnstructuredEmailLoader
from langchain_community.document_loaders import UnstructuredPowerPointLoader
from qanything_kernel.utils.loader import UnstructuredPaddlePDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from qanything_kernel.utils.loader.csv_loader import CSVLoader
from qanything_kernel.utils.loader.markdown_parser import convert_markdown_to_langchaindoc
from qanything_kernel.connector.database.mysql.mysql_client import KnowledgeBaseManager
import asyncio
import aiohttp
import docx2txt
import base64
import pandas as pd
import os
import json
import requests
import threading
import re
import newspaper
import uuid
import traceback
import openpyxl
import shutil
import time
import docx


def get_ocr_result_sync(image_data, api="ocr"):
    try:
        response = requests.post(f"http://{LOCAL_OCR_SERVICE_URL}/" + api, json=image_data, timeout=120)
        response.raise_for_status()  # 如果请求返回了错误状态码，将会抛出异常
        ocr_res = response.text
        ocr_res = json.loads(ocr_res)
        return ocr_res['result']
    except Exception as e:
        insert_logger.warning(f"ocr error: {traceback.format_exc()}")
        return None

def get_pdf_result_sync(file_path):
    try:
        data = {
            'filename': file_path,
            'save_dir': os.path.dirname(file_path)
        }
        headers = {"content-type": "application/json"}
        response = requests.post(f"http://{LOCAL_PDF_PARSER_SERVICE_URL}/pdfparser", json=data, headers=headers,
                                 timeout=240)
        response.raise_for_status()  # 如果请求返回了错误状态码，将会抛出异常
        response_json = response.json()
        markdown_file = response_json.get('markdown_file')
        return markdown_file
    except Exception as e:
        insert_logger.warning(f"pdf parser error: {traceback.format_exc()}")
        return None

def get_pdf_chunk_result_sync(file_path, chunk_size=800):
    try:
        data = {
            'filename': file_path,
            'chunk_size': chunk_size,
            'save_dir': os.path.dirname(file_path)
        }
        insert_logger.info(f"pdf parser : {data}")
        headers = {"content-type": "application/json"}
        response = requests.post(f"http://{LOCAL_PDF_PARSER_SERVICE_URL}/pdfparser_chunk", json=data, headers=headers)
        response.raise_for_status()  # 如果请求返回了错误状态码，将会抛出异常
        response_json = response.json()
        chunks = response_json.get('chunks')
        chunks_document = []
        for chunk in chunks:
            metadata = {key: value for key, value in chunk.items() if key != "text"}
            tmp_chunk = Document(page_content=chunk.get("text"), metadata=metadata)
            chunks_document.append(tmp_chunk)
        return chunks_document
    except Exception as e:
        insert_logger.warning(f"pdf parser error: {traceback.format_exc()}")
        return None

def get_docx_result_sync(docx_path, image_save_dir, markdown_template="![figure]({image_path})"):
    """
    处理Word文档：
    1. 提取所有图片保存到本地
    2. 生成新文本，图片位置用Markdown格式替换
    """
    # 创建保存目录
    import hashlib
    os.makedirs(image_save_dir, exist_ok=True)

    doc = docx.Document(docx_path)
    output_content = []
    img_counter = 0  # 给图片唯一编号
    runs_within_single_paragraph = []

    # 遍历文档所有元素
    for element in doc.element.body:
        # 处理段落中的图片
        if element.tag.endswith('p'):
            para = doc.paragraphs[output_content.count("[PARA_END]")]
            output_content.append("[PARA_END]")  # 段落分隔标记

            if para.text.strip():
                runs_within_single_paragraph.append(para.text)
            for run in para.runs:
                # 检测内联图片
                if 'graphic' in run._element.xml:
                    img_counter += 1
                    # 提取图片二进制数据
                    embed_attrs = run._element.xpath('.//a:blip/@r:embed')
                    if not embed_attrs:
                        insert_logger.warning("跳过无效图片")
                        continue  # 跳过无效图片
                        
                    img_blob = embed_attrs[0]
                    image_part = doc.part.related_parts[img_blob]
                    image_data = image_part.blob

                    # 生成唯一文件名（哈希+索引避免重复）
                    file_hash = hashlib.md5(image_data).hexdigest()[:8]
                    file_ext = "jpg" #image_part.content_type.split("/")[-1]
                    filename = f"img_{img_counter}_{file_hash}.{file_ext}"
                    save_path = os.path.join(image_save_dir, filename)

                    # 保存图片
                    with open(save_path, 'wb') as f:
                        f.write(image_data)

                    # 插入Markdown引用
                    runs_within_single_paragraph.append(markdown_template.format(image_path=filename))
        
        # 处理表格中的图片（根据需要扩展）
        elif element.tag.endswith('tbl'):
            tb = doc.tables[output_content.count("[TABLE_CONTENT]")]
            output_content.append("[TABLE_CONTENT]")  # 示例占位符
            
            from docx.table import Table
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def get_colspan(cell):
                """获取单元格的跨列数"""
                tc = cell._tc
                grid_span_elem = tc.tcPr.find(qn('w:gridSpan'))
                if grid_span_elem is not None:
                    return int(grid_span_elem.get(qn('w:val'), 1))
                return 1

            def get_rowspan(table, row_idx, col_idx, max_cols):
                """获取单元格的跨行数"""
                if col_idx >= max_cols:
                    return 1
                try:
                    cell = table.cell(row_idx, col_idx)
                except IndexError:
                    return 1
                
                tc = cell._tc
                v_merge = tc.tcPr.find(qn('w:vMerge'))
                if v_merge is None or v_merge.get(qn('w:val'), "") != "restart":
                    return 1

                rowspan = 1
                current_row = row_idx + 1
                while current_row < len(table.rows):
                    try:
                        next_cell = table.cell(current_row, col_idx)
                    except IndexError:
                        break
                    
                    next_tc = next_cell._tc
                    next_v_merge = next_tc.tcPr.find(qn('w:vMerge'))
                    if next_v_merge is not None:
                        next_val = next_v_merge.get(qn('w:val'), "")
                        if next_val == "restart" and str(next_cell.text) == str(cell.text):
                            rowspan += 1
                            current_row += 1
                        else:
                            break
                    else:
                        break
                return rowspan
            
            def handle_merged_cells(tb: Table):
                """提取表格的合并信息"""
                num_rows = len(tb.rows)
                num_cols = max(len(row.cells) for row in tb.rows)
                spans = [[False for _ in range(num_cols)] for _ in range(num_rows)]
                
                merged_info = []
                for row_idx in range(num_rows):
                    row_info = []
                    col_idx = 0
                    while col_idx < num_cols:
                        if spans[row_idx][col_idx]:
                            col_idx += 1
                            continue
                        
                        try:
                            cell = tb.cell(row_idx, col_idx)
                        except IndexError:
                            col_idx += 1
                            continue

                        colspan = get_colspan(cell)
                        rowspan = get_rowspan(tb, row_idx, col_idx, num_cols)
                        # 确保不超出表格边界
                        colspan = min(colspan, num_cols - col_idx)
                        rowspan = min(rowspan, num_rows - row_idx)

                        # 标记覆盖的位置
                        for r in range(rowspan):
                            for c in range(colspan):
                                if row_idx + r < num_rows and col_idx + c < num_cols:
                                    spans[row_idx + r][col_idx + c] = True
                        
                        row_info.append({
                            'text': cell.text.strip(),
                            'colspan': colspan,
                            'rowspan': rowspan
                        })
                    merged_info.append(row_info)
                return merged_info

            def split_table(tb: Table, max_data_rows=30, max_chars=1000):
                """智能拆分带合并单元格的表格"""
                merged_info = handle_merged_cells(tb)
                sub_tables = []
                current_sub = []
                current_chars = 0
                current_rows = 0

                # 保留表头
                header = merged_info[0]
                current_sub.append(header)
                current_chars += sum(len(c['text']) for c in header)
                current_rows += 1

                while current_rows<len(merged_info):
                    # 计算新行的字符数
                    row_info = merged_info[current_rows]
                    row_chars = sum(len(c['text']) for c in row_info)
                    new_chars = current_chars + row_chars
                    new_rows = current_rows + 1

                    # 检查合并单元格的连续性
                    has_vertical_merge = any(c['rowspan'] >1 for c in row_info)
                    if has_vertical_merge:
                        # 拆分条件：超过行数限制 或 超过字符限制 且 无垂直合并延续
                        if (len(current_sub)+1 > max_data_rows or new_chars > max_chars):
                            sub_tables.append(current_sub)
                            # 重置时包含表头和当前行
                            current_sub = [header, row_info]
                            current_chars = sum(len(c['text']) for c in header) + row_chars
                            current_rows = new_rows
                        else:
                            current_sub.append(row_info)
                            current_chars = new_chars
                            current_rows = new_rows
                                                
                        max_rowspan = max(c['rowspan'] for c in row_info)
                        for i in range(max_rowspan-1):
                            current_sub.append(merged_info[current_rows+i])
                            current_chars = current_chars + sum(len(c['text']) for c in merged_info[current_rows+i])
                        current_rows += max_rowspan - 1
                        continue
                            

                    # 拆分条件：超过行数限制 或 超过字符限制 且 无垂直合并延续
                    if (len(current_sub)+1 > max_data_rows or new_chars > max_chars):
                        sub_tables.append(current_sub)
                        # 重置时包含表头和当前行
                        current_sub = [header, row_info]
                        current_chars = sum(len(c['text']) for c in header) + row_chars
                        current_rows = new_rows
                    else:
                        current_sub.append(row_info)
                        current_chars = new_chars
                        current_rows = new_rows

                # 添加最后的子表
                if current_sub:
                    sub_tables.append(current_sub)
                
                return sub_tables

            def convert_to_html_with_merge(merged_info):
                """将带合并信息的表格转换为HTML"""
                html = ['<table border="0">']
                for row in merged_info:
                    html.append('<tr>')
                    for cell in row:
                        if cell['rowspan'] == 0:
                            # 跳过垂直合并的延续行
                            continue
                        attrs = []
                        if cell['colspan'] > 1:
                            attrs.append(f'colspan="{cell["colspan"]}"')
                        if cell['rowspan'] > 1:
                            attrs.append(f'rowspan="{cell["rowspan"]}"')
                        attr_str = ' '.join(attrs)
                        html.append(f'<td {attr_str}>{cell["text"]}</td>')
                    html.append('</tr>')
                html.append('</table>')
                return ''.join(html)
            
            # 判断是否需要拆分，超过行数限制或者字数限制均进行拆分
            tbl_line_limit = 30
            tbl_char_limit = 1000
            sub_tables = split_table(tb, tbl_line_limit, tbl_char_limit)
            for sub in sub_tables:
                # 生成带合并属性的HTML
                html_table = convert_to_html_with_merge(sub)
                runs_within_single_paragraph.append(html_table)
            

    new_docs = []
    current_clause = []
    clauses = []
    CLAUSE_PATTERN = re.compile(r'^第\S+条')  # 匹配条款标题
    current_length = 0
    MAX_LENGTH = 5000  # 最大长度限制
    TABLE_FLAG = "</table>"  # 表格标识
    
    # 预处理队列, 构建可回溯的队列
    processing_queue = []    
    for idx, content in enumerate(runs_within_single_paragraph):
        processing_queue.append({
            "index": idx,
            "content": content.strip(),
            "processed": False
        })

    idx = 0
    while idx < len(processing_queue):
        item = processing_queue[idx]
        if item["processed"]:
            idx += 1
            continue

        if item["content"]=="[chunk-split]":
            item["processed"] = True
            if current_clause:
                clauses.append("\n".join(current_clause))
                current_clause = []
                current_length = 0
            idx += 1
            continue
            
        content = item["content"]
        content_len = len(content)
        
        # 条款处理逻辑（优先级最高）
        if CLAUSE_PATTERN.match(content):
            # 保存已有内容
            if current_clause:
                clauses.append("\n".join(current_clause))
                current_clause = []
                current_length = 0
            # 处理新条款
            current_clause.append(content)
            current_length = content_len
            item["processed"] = True
            idx += 1
            # 长度检查
            if current_length >= MAX_LENGTH:
                clauses.append("\n".join(current_clause))
                current_clause = []
                current_length = 0
            continue
            
        # 表格处理逻辑（第二优先级）
        if TABLE_FLAG in content:
            table_group = []
            pattern = r'^(.*表|表.*)$'

            # 向前查找表格描述
            prev_idx = idx - 1
            if prev_idx >= 0: # and not processing_queue[prev_idx]["processed"]:
                prev_content = processing_queue[prev_idx]["content"]
                if re.match(pattern, prev_content):
                    table_group.insert(0, prev_content)
                    processing_queue[prev_idx]["processed"] = True
            # 添加当前表格
            table_group.append(content)
            item["processed"] = True
            
            # 向后查找表格描述
            next_idx = idx + 1
            if next_idx < len(processing_queue):
                next_content = processing_queue[next_idx]["content"]
                if re.match(pattern, next_content):
                    table_group.append(next_content)
                    processing_queue[next_idx]["processed"] = True
            
            # 保存已有内容
            if current_clause:
                clauses.append("\n".join(current_clause))
                current_clause = []
                current_length = 0
            # 合并表格内容，将表格作为独立条款
            clauses.append("\n".join(table_group))
            
            idx = next_idx  # 跳转到下一个未处理位置
            continue
            
        # 普通文本处理
        if current_length + content_len > MAX_LENGTH:
            clauses.append("\n".join(current_clause))
            current_clause = [content]
            current_length = content_len
        else:
            current_clause.append(content)
            current_length += content_len
        
        item["processed"] = True
        idx += 1

    # 处理剩余内容
    if current_clause:
        clauses.append("\n".join(current_clause))

    for clause in clauses:
        new_doc = Document(page_content=clause)
        new_doc.metadata["has_table"] = True if "</table>" in clause else False
        new_doc.metadata["images"] = re.findall(r'!\[figure\]\(.*?\)', clause)
        new_docs.append(new_doc)

    return new_docs

def get_html_table_result_sync(xlsx_path, sheet_name, tbl_line_limit=100, tbl_char_limit=4000):
    """
    支持合并单元格的Excel转HTML切分方案
    
    参数说明与之前一致...
    """
    # 使用openpyxl读取合并单元格信息
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name]
    
    # 获取所有合并区域
    merged_ranges = []
    for merged_cell in ws.merged_cells.ranges:
        merged_ranges.append({
            'min_row': merged_cell.min_row,
            'max_row': merged_cell.max_row,
            'min_col': merged_cell.min_col,
            'max_col': merged_cell.max_col
        })
    
    # 构建单元格属性矩阵
    data = []
    for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
        html_row = []
        for col_idx, cell_value in enumerate(row, 1):
            # 检查是否属于合并区域
            is_merged = False
            for mr in merged_ranges:
                if (row_idx >= mr['min_row'] and row_idx <= mr['max_row'] and 
                    col_idx >= mr['min_col'] and col_idx <= mr['max_col']):
                    if row_idx == mr['min_row'] and col_idx == mr['min_col']:
                        html_row.append({
                            'value': cell_value,
                            'rowspan': mr['max_row'] - mr['min_row'] + 1,
                            'colspan': mr['max_col'] - mr['min_col'] + 1
                        })
                    is_merged = True
                    break
            if not is_merged:
                html_row.append({'value': cell_value, 'rowspan': 1, 'colspan': 1})
        data.append(html_row)
    
    # 转换为HTML的函数
    def chunk_to_html(chunk):
        html = ['<table border="1">']
        for row in chunk:
            html.append('<tr>')
            for cell in row:
                if cell:  # 跳过被合并的单元格
                    html.append(
                        f'<td rowspan="{cell["rowspan"]}" colspan="{cell["colspan"]}">'
                        f'{cell["value"] if cell["value"] is not None else ""}</td>'
                    )
            html.append('</tr>')
        html.append('</table>')
        return ''.join(html)
    
    # 切分逻辑（需要重新处理）
    result = []
    current_chunk = []
    current_char_count = 0
    
    for row in data:
        # 预计算当前行字符数（近似值）
        row_char = sum(len(str(cell['value'])) for cell in row)
        
        # 检查是否超过限制
        if (len(current_chunk) + 1 > tbl_line_limit or 
            current_char_count + row_char > tbl_char_limit):
            
            # 处理合并单元格跨越切分的问题
            has_cross_merge = any(
                cell['rowspan'] > 1 and (len(current_chunk)+1) >= cell['rowspan']
                for cell in current_chunk[-1] if cell
            ) if current_chunk else False
            
            if has_cross_merge and current_chunk:
                # 强制包含跨越行
                split_row = len(current_chunk) - 1
                result.append(chunk_to_html(current_chunk[:split_row]))
                current_chunk = current_chunk[split_row:]
                current_char_count = sum(len(str(cell['value'])) for r in current_chunk for cell in r)
            
            else:
                # 正常切分
                if current_chunk:
                    result.append(chunk_to_html(current_chunk))
                    current_chunk = []
                    current_char_count = 0
        
        # 添加新行
        current_chunk.append(row)
        current_char_count += row_char
    
    # 处理剩余部分
    if current_chunk:
        result.append(chunk_to_html(current_chunk))
    
    docs = []
    for page_content in result:
        new_doc = Document(page_content=page_content)
        new_doc.metadata["has_table"] = True if "</table>" in page_content else False
        docs.append(new_doc)

    return docs

def get_xlsx_result_sync(file_path, sheet_name):
    xlsx = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl', header=None)
    xlsx = xlsx.dropna(how='all', axis=1)  # 只删除全为空的列
    xlsx = xlsx.dropna(how='all', axis=0)  # 只删除全为空的行

    
    title = ""
    # 循环处理可能的标题行（例如跨多行合并的标题）,判断开始的行，判断是否为标题行
    while not xlsx.empty:
        current_row = list(xlsx.iloc[0])
        # 检查第一列非空且其他列全为空
        if not pd.isna(current_row[0]) and all(pd.isna(cell) for cell in current_row[1:]):
            title += str(current_row[0])+'\n'
            xlsx = xlsx.iloc[1:].reset_index(drop=True)  # 移除已处理的行并重置索引
        else:
            break
    if title:
        insert_logger.info(f'{file_path}.{sheet_name} found tiltle: {title}')

    # 判断header行，如果没有列合并单元格，则直接取第一行，认定为简单表格，否则则认定为复杂表格，直接将表格按html进行处理
    end_header_index = 0
    row_count = xlsx.shape[0]   # 获取行数
    while end_header_index< row_count:
        current_row = list(xlsx.iloc[end_header_index])
        insert_logger.debug(end_header_index, current_row)
        end_header_index+=1
        if all(not pd.isna(cell) for cell in current_row):
            break
    
    docs = []
    if end_header_index == 1:   # 简单表格进行处理
        insert_logger.info(f'{file_path}.{sheet_name} is simple table')
        header_rows = xlsx.iloc[:end_header_index].fillna('')
        columns = [' '.join(col).strip() for col in header_rows.transpose().values]
        xlsx.columns = columns
        xlsx = xlsx.iloc[end_header_index:].reset_index(drop=True)  # 移除已处理的行并重置索引

        csv_file_path = file_path[:-5] + f'_{sheet_name}.csv'
        xlsx.to_csv(csv_file_path, index=False)
        insert_logger.info('xlsx2csv: %s', csv_file_path)
        loader = CSVLoader(csv_file_path, autodetect_encoding=True,
                            csv_args={"delimiter": ",", "quotechar": '"'})
        docs.extend(loader.load())

        # xlsx_transposed = xlsx.transpose()
        # csv_file_transposed_path = file_path[:-5] + f'_{sheet_name}_transposed.csv'
        # xlsx_transposed.to_csv(csv_file_transposed_path, index=False)
        # insert_logger.info('xlsx2csv: %s', csv_file_transposed_path)
        # loader_transposed = CSVLoader(csv_file_transposed_path, autodetect_encoding=True,
        #                     csv_args={"delimiter": ",", "quotechar": '"'})
        # docs.extend(loader_transposed.load())
        if title:
            for doc in docs:
                doc.metadata["title_lst"] = [title]
                doc.page_content = '##'+ title + doc.page_content
    elif not xlsx.empty:   # 复杂表格进行处理
        insert_logger.info(f'{file_path}.{sheet_name} is complex table')
        docs = get_html_table_result_sync(file_path, sheet_name)
        if title:
            for doc in docs:
                doc.metadata["title_lst"] = [title]
                doc.page_content = '##'+ title + doc.page_content
    else:
        insert_logger.info(f'{file_path}.{sheet_name} just title')
        docs = [Document(page_content=title)]

    # 所有表格的按行进行切分，复杂的自动合并切分，不再进行父切分
    for doc in docs:
        doc.metadata["single_parent"] = True

    return docs



class LocalFileForInsert:
    def __init__(self, user_id, kb_id, file_id, file_location, file_name, file_url, chunk_size, mysql_client):
        self.chunk_size = chunk_size
        self.markdown_text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=0,
                                                                     length_function=num_tokens_embed)
        self.user_id = user_id
        self.kb_id = kb_id
        self.file_id = file_id
        self.docs: List[Document] = []
        self.embs = []
        self.file_name = file_name
        self.file_location = file_location
        self.file_url = ""
        self.faq_dict = {}
        self.file_path = ""
        self.mysql_client = mysql_client
        self.file_content_text = ""
        self.outline = ""
        self.summary = ""
        if self.file_location == 'FAQ':
            faq_info = self.mysql_client.get_faq(self.file_id)
            user_id, kb_id, question, answer, nos_keys = faq_info
            self.faq_dict = {'question': question, 'answer': answer, 'nos_keys': nos_keys}
        elif self.file_location == 'URL':
            self.file_url = file_url
            upload_path = os.path.join(UPLOAD_ROOT_PATH, user_id)
            file_dir = os.path.join(upload_path, self.kb_id, self.file_id)
            os.makedirs(file_dir, exist_ok=True)
            self.file_path = os.path.join(file_dir, self.file_name)
        else:
            self.file_path = self.file_location
        self.event = threading.Event()
        self.mysql_client = KnowledgeBaseManager()
        self.separators = self.mysql_client.get_kb_parser_config(self.kb_id)["separators"]

    @staticmethod
    @get_time
    def image_ocr_txt(filepath, dir_path="tmp_files"):
        full_dir_path = os.path.join(os.path.dirname(filepath), dir_path)
        if not os.path.exists(full_dir_path):
            os.makedirs(full_dir_path)
        filename = os.path.split(filepath)[-1]

        # 读取图片
        img_np = open(filepath, 'rb').read()

        img_data = {
            "img64": base64.b64encode(img_np).decode("utf-8"),
        }

        result = get_ocr_result_sync(img_data)
        insert_logger.info(f"ocr result: {result}")

        ocr_result = [line for line in result if line]
        ocr_result = '\n'.join(ocr_result)

        insert_logger.info(f'ocr_res[:100]: {ocr_result[:100]}')

        # 写入结果到文本文件
        txt_file_path = os.path.join(full_dir_path, "%s.txt" % (filename))
        with open(txt_file_path, 'w', encoding='utf-8') as fout:
            fout.write(ocr_result)

        return txt_file_path

    def table_process(self, doc):
        table_infos = get_table_infos(doc.page_content)
        title_lst = doc.metadata['title_lst']
        new_docs = []
        if table_infos is not None:
            tmp_content = '\n'.join(title_lst) + '\n' + doc.page_content
            if num_tokens_embed(tmp_content) <= self.chunk_size:
                doc.page_content = tmp_content
                return [doc]
            head_line = table_infos['head_line']
            end_line = table_infos['end_line']
            table_head = table_infos['head']

            # 处理表格前的内容
            if head_line != 0:
                tmp_doc = Document(
                    page_content='\n'.join(title_lst) + '\n' + '\n'.join(table_infos['lines'][:head_line]),
                    metadata=doc.metadata)
                new_docs.append(tmp_doc)

            # 处理表格内容
            table_content = '\n'.join(title_lst) + '\n' + table_head
            for line in table_infos['lines'][head_line + 2:end_line + 1]:
                if num_tokens_embed(table_content + '\n' + line) > self.chunk_size:
                    # 如果添加新行会超出chunk_size，先保存当前内容
                    tmp_doc = Document(page_content=table_content, metadata=doc.metadata)
                    new_docs.append(tmp_doc)
                    # 重新开始一个新的chunk，包含标题和表头
                    table_content = '\n'.join(title_lst) + '\n' + table_head + '\n' + line
                else:
                    if line == table_head.split('\n')[0]:
                        table_content += '\n\n' + line
                        # print('match table_head:', table_content)
                    else:
                        table_content += '\n' + line

            # 保存最后一个chunk
            if table_content != '\n'.join(title_lst) + '\n' + table_head:
                tmp_doc = Document(page_content=table_content, metadata=doc.metadata)
                new_docs.append(tmp_doc)

            # 处理表格后的内容
            if end_line != len(table_infos['lines']) - 1:
                tmp_doc = Document(
                    page_content='\n'.join(title_lst) + '\n' + '\n'.join(table_infos['lines'][end_line:]),
                    metadata=doc.metadata)
                new_docs.append(tmp_doc)

            insert_logger.info(f"TABLE SLICES: {new_docs[:2]}")
        else:
            return None
        return new_docs

    @staticmethod
    def get_page_id(doc, pre_page_id):
        # 查找 page_id 标志行
        lines = doc.page_content.split('\n')
        for line in lines:
            if re.match(r'^#+ 当前页数:\d+$', line):
                try:
                    page_id = int(line.split(':')[-1])
                    return page_id
                except ValueError:
                    continue
        return pre_page_id

    def markdown_process(self, docs: List[Document]):
        new_docs = []
        for doc in docs:
            title_lst = doc.metadata['title_lst']
            # 删除所有仅有多个#的title
            title_lst = [t for t in title_lst if t.replace('#', '') != '']
            has_table = doc.metadata['has_table']
            if has_table:
                table_doc_id = str(uuid.uuid4())
                self.mysql_client.add_document(table_doc_id, doc.to_json())
                doc.metadata['table_doc_id'] = table_doc_id
                table_docs = self.table_process(doc)
                if table_docs:
                    new_docs.extend(table_docs)
                    continue
            slices = self.markdown_text_splitter.split_documents([doc])
            # insert_logger.info(f"markdown_text_splitter: {len(slices)}")
            if len(slices) == 1:
                slices[0].page_content = '\n\n'.join(title_lst) + '\n\n' + slices[0].page_content
            else:
                for idx, slice in enumerate(slices):
                    slice.page_content = '\n\n'.join(
                        title_lst) + f'\n\n###### 第{idx + 1}段内容如下：\n' + slice.page_content
            new_docs.extend(slices)
        return new_docs

    @get_time_async
    async def url_to_documents_async(self, file_path, file_name, file_url, dir_path="tmp_files", max_retries=3):
        full_dir_path = os.path.join(os.path.dirname(file_path), dir_path)
        if not os.path.exists(full_dir_path):
            os.makedirs(full_dir_path)

        for attempt in range(max_retries):
            try:
                headers = {
                    "Accept": "application/json",
                    "X-Return-Format": "markdown",
                    "X-Timeout": "15",
                }
                async with aiohttp.ClientSession() as session:
                    async with session.get(f"https://r.jina.ai/{file_url}", headers=headers, timeout=30) as response:
                        jina_response = await response.json()
                        if jina_response['code'] == 200:
                            title = jina_response['data'].get('title', '')
                            markdown_str = jina_response['data'].get('content', '')
                            markdown_str = html_to_markdown(markdown_str)
                            md_file_path = os.path.join(full_dir_path, "%s.md" % (file_name))
                            with open(md_file_path, 'w', encoding='utf-8') as fout:
                                fout.write(markdown_str)
                            docs = convert_markdown_to_langchaindoc(md_file_path)
                            if title:
                                for doc in docs:
                                    doc.metadata['title'] = title
                            docs = self.markdown_process(docs)
                            return docs
                        else:
                            insert_logger.warning(f"jina get url warning: {file_url}, {jina_response}")
            except Exception as e:
                insert_logger.warning(f"jina get url error: {file_url}, {traceback.format_exc()}")

            if attempt < max_retries - 1:  # 如果不是最后一次尝试，等待30秒后重试
                await asyncio.sleep(30)

        return None

    @get_time
    def url_to_documents(self, file_path, file_name, file_url, dir_path="tmp_files", max_retries=3):
        full_dir_path = os.path.join(os.path.dirname(file_path), dir_path)
        if not os.path.exists(full_dir_path):
            os.makedirs(full_dir_path)

        for attempt in range(max_retries):
            try:
                headers = {
                    "Accept": "application/json",
                    "X-Return-Format": "markdown",
                    "X-Timeout": "15",
                }
                response = requests.get(f"https://r.jina.ai/{file_url}", headers=headers, timeout=30)
                jina_response = response.json()
                if jina_response['code'] == 200:
                    title = jina_response['data'].get('title', '')
                    markdown_str = jina_response['data'].get('content', '')
                    markdown_str = html_to_markdown(markdown_str)
                    md_file_path = os.path.join(full_dir_path, "%s.md" % (file_name))
                    with open(md_file_path, 'w', encoding='utf-8') as fout:
                        fout.write(markdown_str)
                    docs = convert_markdown_to_langchaindoc(md_file_path)
                    if title:
                        for doc in docs:
                            doc.metadata['title'] = title
                    docs = self.markdown_process(docs)
                    return docs
                else:
                    insert_logger.warning(f"jina get url warning: {file_url}, {jina_response}")
            except Exception as e:
                insert_logger.warning(f"jina get url error: {file_url}, {traceback.format_exc()}")

            if attempt < max_retries - 1:  # 如果不是最后一次尝试，等待30秒后重试
                time.sleep(30)

        return None

    @staticmethod
    def excel_to_markdown(file_path, markdown_path):
        def clean_cell_content(cell):
            if cell is None:
                return ''
            # 将单元格内容转换为字符串，并替换换行符为空格
            return re.sub(r'\s+', ' ', str(cell)).strip()
        basename = os.path.splitext(os.path.basename(file_path))[0]
        markdown_file = os.path.join(markdown_path, f"{basename}.md")

        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        with open(markdown_file, 'w', encoding='utf-8') as md_file:
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                md_file.write(f"# {sheet_name}\n\n")

                # 获取非空行和列
                rows = [[clean_cell_content(cell) for cell in row] for row in sheet.iter_rows(values_only=True)]
                non_empty_rows = [row for row in rows if any(cell != '' for cell in row)]

                if not non_empty_rows:
                    continue  # 跳过空表格

                # 判断第一行，是否为标题行
                first_row = non_empty_rows[0][1:]
                if all(cell == '' for cell in first_row):
                    insert_logger.info("第一行是标题行")
                    md_file.write(f"## {non_empty_rows[0][0]}\n\n")
                    non_empty_rows = non_empty_rows[1:]

                max_cols = max(len(row) for row in non_empty_rows)

                # 处理每一行
                for row_index, row in enumerate(non_empty_rows):
                    # 补齐空单元格
                    padded_row = row + [''] * (max_cols - len(row))

                    # 转换为Markdown表格行，使用竖线作为分隔符
                    markdown_row = '| ' + ' | '.join(padded_row) + ' |'
                    md_file.write(markdown_row + '\n')

                    # 在第一行后添加分隔符
                    if row_index == 0:
                        separator = '|' + '|'.join(['---' for _ in range(max_cols)]) + '|'
                        md_file.write(separator + '\n')

                md_file.write('\n\n')  # 在每个表格后添加空行

        insert_logger.info(f"转换完成。Markdown 文件已保存为 {markdown_file}")
        return markdown_file

    @staticmethod
    def load_text(file_path):
        encodings = ['utf-8', 'iso-8859-1', 'windows-1252']

        for encoding in encodings:
            try:
                loader = TextLoader(file_path, encoding=encoding)
                docs = loader.load()
                insert_logger.info(f"TextLoader {encoding} success: {file_path}")
                return docs
            except Exception:
                insert_logger.warning(f"TextLoader {encoding} error: {file_path}, {traceback.format_exc()}")

        insert_logger.error(f"Failed to load file with all attempted encodings: {file_path}")
        return []

    @staticmethod
    def copy_images(image_root_path, output_dir):
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        # 获取当前目录下所有jpg文件
        images = [f for f in os.listdir(image_root_path) if f.endswith('.jpg')]
        # 复制到指定目录
        for image in images:
            single_image_path = os.path.join(image_root_path, image)
            insert_logger.info(f"copy image: {single_image_path} -> {output_dir}")
            shutil.copy(single_image_path, output_dir)

    @get_time
    def split_file_to_docs(self):
        insert_logger.info(f"start split file to docs, file_path: {self.file_name}")
        if self.faq_dict:
            docs = [Document(page_content=self.faq_dict['question'], metadata={"faq_dict": self.faq_dict})]
        elif self.file_url:
            insert_logger.info("load url: {}".format(self.file_url))
            docs = self.url_to_documents(self.file_path, self.file_name, self.file_url)
            if docs is None:
                try:
                    article = newspaper.article(self.file_url, timeout=120)
                    docs = [
                        Document(page_content=article.text, metadata={"title": article.title, "url": self.file_url})]
                except Exception as e:
                    insert_logger.error(f"newspaper get url error: {self.file_url}, {traceback.format_exc()}")
                    loader = MyRecursiveUrlLoader(url=self.file_url)
                    docs = loader.load()
        elif self.file_path.lower().endswith(".md"):
            try:
                docs = convert_markdown_to_langchaindoc(self.file_path)
                docs = self.markdown_process(docs)
            except Exception as e:
                insert_logger.error(
                    f"convert_markdown_to_langchaindoc error: {self.file_path}, {traceback.format_exc()}")
                loader = UnstructuredFileLoader(self.file_path, strategy="fast")
                docs = loader.load()
        elif self.file_path.lower().endswith(".txt"):
            page_content = open(self.file_path, 'r', encoding='utf-8').read()
            chunk_state = "[chunk-split]" in page_content
            page_chunks = page_content.split("\n[chunk-split]\n")
            docs = []
            for page_chunk in page_chunks:
                if page_chunk:
                    docs.append(Document(page_content=page_chunk, metadata={"single_parent":chunk_state}))
                    

        elif self.file_path.lower().endswith(".pdf"):
            # markdown_file = get_pdf_result_sync(self.file_path)
            # if markdown_file:
            #     docs = convert_markdown_to_langchaindoc(markdown_file)
            #     docs = self.markdown_process(docs)
            #     images_dir = os.path.join(IMAGES_ROOT_PATH, self.file_id)
            #     self.copy_images(os.path.dirname(markdown_file), images_dir)
            # else:
            #     insert_logger.warning(
            #         f'Error in Powerful PDF parsing, use fast PDF parser instead.')
            #     loader = UnstructuredPaddlePDFLoader(self.file_path, strategy="fast")
            #     docs = loader.load()
            
            insert_logger.info("use pdf chunk")
            docs = get_pdf_chunk_result_sync(self.file_path, self.chunk_size)
            if docs is None:
                insert_logger.warning(
                    f'Error in Powerful PDF parsing, use fast PDF parser instead.')
                loader = UnstructuredPaddlePDFLoader(self.file_path, strategy="fast")
                docs = loader.load()
            images_dir = os.path.join(IMAGES_ROOT_PATH, self.file_id)
            self.copy_images(os.path.dirname(self.file_path), images_dir)
            insert_logger.info(f"pdf load docs sucess")
        elif self.file_path.lower().endswith(".jpg") or self.file_path.lower().endswith(
                ".png") or self.file_path.lower().endswith(".jpeg"):
            txt_file_path = self.image_ocr_txt(filepath=self.file_path)
            loader = TextLoader(txt_file_path, autodetect_encoding=True)
            docs = loader.load()
        elif self.file_path.lower().endswith(".docx"):
            try:
                save_image_path = os.path.dirname(self.file_path)
                docs = get_docx_result_sync(self.file_path, save_image_path)
                self.copy_images(save_image_path, os.path.join(IMAGES_ROOT_PATH, self.file_id))
            except Exception as e:
                insert_logger.warning('Error in Powerful Word parsing, use docx2txt instead.')
                text = docx2txt.process(self.file_path)
                docs = [Document(page_content=text)]
        elif self.file_path.lower().endswith(".xlsx"):
            try:
                docs = []
                excel_file = pd.ExcelFile(self.file_path)
                sheet_names = excel_file.sheet_names
                for idx, sheet_name in enumerate(sheet_names):
                    docs.extend(get_xlsx_result_sync(self.file_path, sheet_name))
            except Exception as e:
                insert_logger.warning(f'Error in Powerful Excel parsing, {self.file_path}')
                insert_logger.error(e)
                
        elif self.file_path.lower().endswith(".pptx"):
            loader = UnstructuredPowerPointLoader(self.file_path, strategy="fast")
            docs = loader.load()
        elif self.file_path.lower().endswith(".eml"):
            loader = UnstructuredEmailLoader(self.file_path, strategy="fast")
            docs = loader.load()
        elif self.file_path.lower().endswith(".csv"):
            loader = CSVLoader(self.file_path, autodetect_encoding=True, csv_args={"delimiter": ",", "quotechar": '"'})
            docs = loader.load()
        else:
            raise TypeError("文件类型不支持，目前仅支持：[md,txt,pdf,jpg,png,jpeg,docx,xlsx,pptx,eml,csv]")
        self.inject_metadata(docs)

    def inject_metadata(self, docs: List[Document]):
        # 这里给每个docs片段的metadata里注入file_id
        new_docs = []
        page_id = 1
        for doc in docs:
            page_content = re.sub(r'\t+', ' ', doc.page_content)  # 将制表符替换为单个空格
            page_content = re.sub(r'\n{3,}', '\n\n', page_content)  # 将三个或更多换行符替换为两个
            page_content = page_content.strip()  # 去除首尾空白字符
            
            # 从page_content中判断是否有页码信息内容，"\"\"\"QAnythingPage{" +str(i+1)+ "}\"\"\""，如果有则删除
            # insert_logger.info(f"测试检索debugger: {page_content}")
            if not doc.metadata.get("page_id"):
                tmp_page_id = page_id
                try:
                    pattern = re.compile(r'"""QAnythingPage{\s*(\d+)\s*}"""')
                    matches = pattern.findall(page_content)
                    if matches:
                        # insert_logger.info(f"检索的page_id为：{page_id}，提取的page id为：{matches}")
                        if page_id!=int(matches[0]):
                            tmp_page_id = int(matches[0])-1
                        page_id = int(matches[-1])+1
                        page_content = pattern.sub('', page_content)
                    # else:
                        # insert_logger.info(f"检索的page_id为：{page_id}，当前page_content没有页码信息内容")
                except Exception as e:
                    insert_logger.error("Page id not match, please check the markdown file.")
                    insert_logger.error(f"error message: {e}")
                    
                if self.file_path.lower().endswith(".pdf"):
                    doc.metadata["page_id"] = tmp_page_id
                


            
            new_doc = Document(page_content=page_content)
            new_doc.metadata["user_id"] = self.user_id
            new_doc.metadata["kb_id"] = self.kb_id
            new_doc.metadata["file_id"] = self.file_id
            new_doc.metadata["file_name"] = self.file_name
            new_doc.metadata["nos_key"] = self.file_location
            new_doc.metadata["file_url"] = self.file_url
            new_doc.metadata["title_lst"] = doc.metadata.get("title_lst", [])
            new_doc.metadata["has_table"] = doc.metadata.get("has_table", False)
            new_doc.metadata["images"] = doc.metadata.get("images", re.findall(r'!\[figure]\(\d+-figure-\d+.jpg.*?\)', page_content))   # 从文本中提取图片数量：![figure]（x-figure-x.jpg）
            new_doc.metadata["page_id"] = doc.metadata.get("page_id", 0)
            new_doc.metadata["bboxes"] = doc.metadata.get("bboxes", [])
            new_doc.metadata["single_parent"] = doc.metadata.get("single_parent", False)
            kb_name = self.mysql_client.get_knowledge_base_name([self.kb_id])[0][2]
            metadata_infos = {"知识库名": kb_name, '文件名': self.file_name}
            new_doc.metadata['headers'] = metadata_infos
            new_doc.metadata['faq_dict'] = doc.metadata.get('faq_dict', {})    
            new_docs.append(new_doc)
        
        if new_docs:
            insert_logger.info('langchain analysis content head: %s', new_docs[0].page_content[:100])
        else:
            insert_logger.info('langchain analysis docs is empty!')
        
        # merge short docs
        insert_logger.info(f"before merge doc lens: {len(new_docs)}")
        child_chunk_size = int(self.chunk_size / 2)
        merged_docs = []
        for doc_idx, doc in enumerate(new_docs):
            if doc.metadata["single_parent"] or not merged_docs:
                merged_docs.append(doc)
            else:
                last_doc = merged_docs[-1]
                # insert_logger.info(f"doc_idx: {doc_idx}, doc_content: {doc.page_content[:100]}")
                # insert_logger.info(f"last_doc_len: {num_tokens_embed(last_doc.page_content)}, doc_len: {num_tokens_embed(doc.page_content)}")
                if num_tokens_embed(last_doc.page_content) + num_tokens_embed(doc.page_content) <= child_chunk_size or doc.page_content.startswith('[figure]'): # or num_tokens_embed(doc.page_content) < child_chunk_size / 4:
                    tmp_content_slices = doc.page_content.split('\n')   # 从第二个chunk开始按段落拆分
                    # print(last_doc.metadata['title_lst'], tmp_content)
                    tmp_content_slices_clear = [line for line in tmp_content_slices if clear_string(line) not in
                                                [clear_string(t) for t in last_doc.metadata['title_lst']]]  # 去掉文本中的标题
                    tmp_content = '\n'.join(tmp_content_slices_clear)
                    # for title in last_doc.metadata['title_lst']:
                    #     tmp_content = tmp_content.replace(title, '')
                    last_doc.page_content += '\n\n' + tmp_content
                    # for title in last_doc.metadata['title_lst']:
                    #     last_doc.page_content = self.remove_substring_after_first(last_doc.page_content, '![figure]')
                    last_doc.metadata['title_lst'] += doc.metadata.get('title_lst', [])
                    last_doc.metadata['has_table'] = last_doc.metadata.get('has_table', False) or doc.metadata.get(
                        'has_table', False)
                    last_doc.metadata['images'] += doc.metadata.get('images', [])
                else:
                    merged_docs.append(doc)
        insert_logger.info(f"after merge doc lens: {len(merged_docs)}")
        self.docs = merged_docs


    def parser_outline(self):
        
        # 直接用python-docx读取文件内容大纲
        toc = ""
        if self.file_path.lower().endswith(".docx"):
            toc = get_docx_toc(self.file_path)
        if len(toc) > 50:
            insert_logger.info(f"直接获取到文件目录内容，长度为{len(toc)}，直接使用该内容作为大纲。\n{toc}")
            self.outline = toc
            return

        if not self.file_content_text:
            for doc in self.docs:
                self.file_content_text += doc.page_content.strip()+'\n'
        
        # 调用大模型进行大纲提取
        system_prompt = OUTLINE_EXTRACT_SYSTEM_PROMPT

        prompt = f"请对输入的文本内容进行大纲提取，并输出大纲内容。输入文本内容为：{self.file_content_text}\n\n大纲内容:"

        def call_llm_api(input_query):
            try:
                from openai import OpenAI
                llm_api = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
                completion = llm_api.chat.completions.create(
                    model = LLM_MODEL_NAME,
                    messages=[
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': input_query},
                    ],
                    temperature=LLM_TEMPERATURE,
                    top_p=LLM_TOP_P
                )

                return remove_think_tags(completion.choices[0].message.content)
            except Exception as e:
                insert_logger.error(f"调用大模型失败: {e}")
                return None


        file_content_tokens = num_tokens_embed(system_prompt+prompt)
        if file_content_tokens > LLM_MAX_LENGTH:
            insert_logger.warning(f"文件内容过长({file_content_tokens} tokens)，超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens)，请缩短内容或分段处理。")
            
            tmp_merge_tokens = 0
            tmp_merge_txt = ""
            chunk_summary = []
            sections = self.file_content_text.split('\n')  # 按照段落进行切分
            for section in sections:
                doc_content_tokens = num_tokens_embed(section)
                if tmp_merge_tokens + doc_content_tokens > LLM_MAX_LENGTH-400:
                    insert_logger.info(f"当前文档内容长度({tmp_merge_tokens}+{doc_content_tokens} tokens)超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens，预留400 tokens用于提示词处理)，进行分段处理。")
                    # 进行分段处理
                    chunk_summary.append(call_llm_api(tmp_merge_txt))
                    tmp_merge_tokens = doc_content_tokens
                    tmp_merge_txt = section+'\n'
                else:
                    tmp_merge_tokens += doc_content_tokens
                    tmp_merge_txt += section+'\n'
                
            if tmp_merge_txt:
                chunk_summary.append(call_llm_api(tmp_merge_txt))
            
            all_summary = "当前输出文本内容过长，已分段处理，以下是每段内容的大纲，请注意每段内容的顺序和对应关系，请根据需要进行合并和调整。\n\n"
            paragraph_id = 1
            for summary in chunk_summary:
                if num_tokens_embed(all_summary+system_prompt+summary) > LLM_MAX_LENGTH:
                    insert_logger.warning(f"分段处理后的大纲内容长度超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens)，先整合当前大纲内容，再进行分段处理。")
                    tmp_result = call_llm_api(all_summary)
                    paragraph_id = 1
                    all_summary = "当前输出文本内容过长，已分段处理，以下是每段内容的大纲，请注意每段内容的顺序和对应关系，请根据需要进行合并和调整。\n\n" + f"第{paragraph_id}段大纲内容: {tmp_result}\n\n"
                else:
                    all_summary += f"第{paragraph_id}段大纲内容: {summary}\n\n"
                paragraph_id += 1

            
            self.outline = call_llm_api(all_summary)
        else:
            self.outline = call_llm_api(prompt)


    def parser_summary(self):
        if not self.file_content_text:
            for doc in self.docs:
                self.file_content_text += doc.page_content.strip()+'\n'
        
        # 调用大模型进行摘要提取和大纲提取
        system_prompt = SUMMARY_EXTRACT_SYSTEM_PROMPT

        prompt = f"输入文本内容为：{self.file_content_text}\n\n摘要总结:"

        def call_llm_api(input_query):
            from openai import OpenAI
            llm_api = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
            completion = llm_api.chat.completions.create(
                model = LLM_MODEL_NAME,
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': input_query},
                ],
                temperature=LLM_TEMPERATURE,
                top_p=LLM_TOP_P
            )

            return remove_think_tags(completion.choices[0].message.content)


        file_content_tokens = num_tokens_embed(system_prompt+prompt)
        if file_content_tokens > LLM_MAX_LENGTH:
            insert_logger.warning(f"文件内容过长({file_content_tokens} tokens)，超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens)，请缩短内容或分段处理。")
            # 进行按照docs切分进行多段处理，处理完成后再进行合并

            tmp_merge_tokens = 0
            tmp_merge_txt = ""
            chunk_summary = []
            sections = self.file_content_text.split('\n')  # 按照段落进行切分
            for section in sections:
                doc_content_tokens = num_tokens_embed(section)
                if tmp_merge_tokens + doc_content_tokens > LLM_MAX_LENGTH-400:
                    insert_logger.info(f"当前文档内容长度({tmp_merge_tokens}+{doc_content_tokens} tokens)超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens，预留400 tokens用于提示词处理)，进行分段处理。")
                    # 进行分段处理
                    chunk_summary.append(call_llm_api(tmp_merge_txt))
                    tmp_merge_tokens = doc_content_tokens
                    tmp_merge_txt = section+'\n'
                else:
                    tmp_merge_tokens += doc_content_tokens
                    tmp_merge_txt += section+'\n'
                
            if tmp_merge_txt:
                chunk_summary.append(call_llm_api(tmp_merge_txt))
            
            all_summary = "当前输出文本内容过长，已分段处理，以下是每段内容的摘要总结，请注意每段内容的顺序和对应关系，请根据需要进行合并和调整。\n\n"
            paragraph_id = 1
            for summary in chunk_summary:
                if num_tokens_embed(all_summary+system_prompt+summary) > LLM_MAX_LENGTH:
                    insert_logger.warning(f"分段处理后的大纲内容长度超过了模型的最大输入长度({LLM_MAX_LENGTH} tokens)，先整合当前大纲内容，再进行分段处理。")
                    tmp_result = call_llm_api(all_summary)
                    paragraph_id = 1
                    all_summary = "当前输出文本内容过长，已分段处理，以下是每段内容的摘要总结，请注意每段内容的顺序和对应关系，请根据需要进行合并和调整。\n\n" + f"第{paragraph_id}段摘要总结内容: {tmp_result}\n\n"
                else:
                    all_summary += f"第{paragraph_id}段摘要总结内容: {summary}\n\n"
                paragraph_id += 1
            
            self.summary = call_llm_api(all_summary)
        else:
            self.summary = call_llm_api(prompt)